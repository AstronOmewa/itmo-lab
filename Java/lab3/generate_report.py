#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Создаем документ
doc = Document()

# Заголовок
title = doc.add_heading('ИТМО, Лабораторная работа 3', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle = doc.add_heading('Объектно-ориентированное проектирование на примере литературного текста', level=2)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

date_para = doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()

# 1. ВВЕДЕНИЕ
doc.add_heading('1. Введение', level=1)
doc.add_paragraph(
    'Целью лабораторной работы является разработка объектно-ориентированной системы для моделирования '
    'сценария на основе произведения русской литературы (история о Шутило и его приключениях). '
    'Проект демонстрирует применение ключевых принципов ООП: наследование, инкапсуляция, полиморфизм '
    'и обработка исключений.'
)

# 2. АРХИТЕКТУРА СИСТЕМЫ
doc.add_heading('2. Архитектура системы', level=1)

doc.add_heading('2.1. Иерархия Story → EventSequence → Event', level=2)
doc.add_paragraph(
    'Основу архитектуры составляет трёхуровневая иерархия для построения и выполнения сценариев:'
)

items = [
    ('Story (История)', 'Содержит набор EventSequence объектов, представляющих последовательности событий. '
                        'Метод tell() выполняет всю историю поочередно запуская все EventSequence.'),
    ('EventSequence (Последовательность событий)', 'Содержит набор Event объектов, которые логически связаны. '
                                                     'Метод simulate() выполняет все события в последовательности, '
                                                     'может иметь условия и временные метки.'),
    ('Event (Событие)', 'Представляет отдельное событие с субъектом (Human), типом события (EventType), '
                        'временем (Time) и объектом воздействия. Метод happen() выводит информацию о событии.')
]

for name, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(name + ': ').bold = True
    p.add_run(desc)

doc.add_heading('2.2. Иерархия типов Item', level=2)
doc.add_paragraph(
    'Все предметы в системе наследуют от абстрактного класса Item, который обеспечивает базовую '
    'функциональность владения предметом и изменения состояния:'
)

doc.add_paragraph('Item (абстрактный)', style='List Bullet')
doc.add_paragraph('├─ Cloth (абстрактная одежда)', style='List Bullet 2')
doc.add_paragraph('│  ├─ Tie (галстук, может быть надет наизнанку)', style='List Bullet 3')
doc.add_paragraph('│  ├─ Pants (брюки, может быть надета неправильно)', style='List Bullet 3')
doc.add_paragraph('│  ├─ Chulki (чулки, могут быть неправильно расположены)', style='List Bullet 3')
doc.add_paragraph('│  └─ Jacket (куртка, содержит инвентарь/карманы)', style='List Bullet 3')
doc.add_paragraph('├─ Car (автомобиль)', style='List Bullet 2')
doc.add_paragraph('└─ DriverLicense (водительское удостоверение)', style='List Bullet 2')

doc.add_heading('2.3. Иерархия типов Human', level=2)
doc.add_paragraph(
    'Все персонажи наследуют от абстрактного класса Human, который содержит инвентарь и базовые методы:'
)

doc.add_paragraph('Human (абстрактный)', style='List Bullet')
doc.add_paragraph('├─ Shutilo (главный персонаж - быстро спешит и путает вещи)', style='List Bullet 2')
doc.add_paragraph('├─ Korzhik (соседка, обнаруживает потери)', style='List Bullet 2')
doc.add_paragraph('└─ Swisstulkin (персонаж помощник)', style='List Bullet 2')

# 3. СИСТЕМА ИСКЛЮЧЕНИЙ
doc.add_heading('3. Система обработки исключений', level=1)

doc.add_paragraph(
    'Проект использует три типа пользовательских исключений для обработки различных ошибочных ситуаций:'
)

exceptions = [
    ('ClothingMiswearException (проверяемое)', 'Выбрасывается когда одежда надевается неправильно '
                                                '(галстук наизнанку, брюки неправильно и т.д.). '
                                                'Содержит информацию о типе одежды и причине ошибки.'),
    ('ItemLostException (проверяемое)', 'Выбрасывается когда предмет теряется во время сценария. '
                                        'Содержит название предмета и имя предыдущего владельца.'),
    ('InventoryFullException (unchecked)', 'Выбрасывается когда инвентарь переполнен (максимум 5 предметов). '
                                           'Расширяет RuntimeException, так как это внутренняя ошибка приложения.')
]

for exc_type, desc in exceptions:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(exc_type + ': ').bold = True
    p.add_run(desc)

# 4. КЛЮЧЕВЫЕ ИНТЕРФЕЙСЫ
doc.add_heading('4. Ключевые интерфейсы', level=1)

interfaces = [
    ('Wearable', 'Для всех носимых предметов. Методы: wear(), wear(Human) с throws ClothingMiswearException'),
    ('WearableInsideOut', 'Расширяет Wearable. Добавляет метод wearInsideOut() для специальных способов надевания'),
    ('Misswearable', 'Для предметов, которые можно неправильно надеть. Метод misswear() throws ClothingMiswearException'),
    ('Missarangeable', 'Для предметов, которые можно неправильно расположить. Метод misarrange() throws ClothingMiswearException'),
    ('Storable', 'Для предметов, которые можно хранить в инвентаре. Метод store(Item)'),
    ('Breakable', 'Для предметов, которые можно сломать. Метод breakItem() возвращает Event'),
    ('Distractable', 'Для персонажей, которые могут отвлекаться. Метод distract() возвращает Event'),
]

for iface, desc in interfaces:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(iface + ': ').bold = True
    p.add_run(desc)

# 5. РЕАЛИЗАЦИЯ equals, hashCode, toString
doc.add_heading('5. Реализация ключевых методов Object', level=1)

doc.add_paragraph(
    'Все классы (13 всего) реализуют три критических метода для корректной работы объектов:'
)

methods = [
    ('equals(Object)', 'Сравнивает два объекта через их hashCode. '
                       'Два объекта считаются равными, если их хеши совпадают.'),
    ('hashCode()', 'Вычисляется с множителем 31 и УЧИТЫВАЕТ ВСЕ поля класса, '
                   'включая сложные объекты (Human, Inventory и т.д.). '
                   'Это обеспечивает корректность работы с коллекциями и сравнениями.'),
    ('toString()', 'Выводит полное представление объекта со всеми значимыми полями. '
                   'Помогает в отладке и логировании.')
]

for method, desc in methods:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(method + ': ').bold = True
    p.add_run(desc)

# 6. ГЛАВНЫЙ КЛАСС Main
doc.add_heading('6. Главный класс Main', level=1)

doc.add_paragraph(
    'Класс Main конструирует полный сценарий без каких-либо System.out.println в своём коде. '
    'Вся логика построена на архитектуре Story→EventSequence→Event:'
)

doc.add_paragraph(
    '1. Создание персонажей (Shutilo, Korzhik, Swisstulkin)\n'
    '2. Создание предметов (одежда, автомобиль, лицензия, цветочный горшок)\n'
    '3. Создание Story объекта\n'
    '4. Построение 4 EventSequence с логическими сценариями:\n'
    '   • Пробуждение и спешка Shutilo\n'
    '   • Одевание с обработкой исключений\n'
    '   • Выход из дома с поломкой цветочного горшка\n'
    '   • Обнаружение потери Korzhik и Swisstulkin\n'
    '5. Добавление всех последовательностей в Story\n'
    '6. Вызов tale.tell() для выполнения всей истории'
)

# 7. МОДЕЛЬ ДАННЫХ
doc.add_heading('7. Модель данных Inventory', level=1)

doc.add_paragraph(
    'Система использует специальную модель для управления инвентарем:'
)

doc.add_paragraph(
    'record InventoryCell: Запись (record) Java 16+ которая хранит Item и его Places (расположение)',
    style='List Bullet'
)

doc.add_paragraph(
    'class Inventory: Управляет коллекцией InventoryCell с ограничением MAX_CAPACITY = 5. '
    'Реализует Storable интерфейс.',
    style='List Bullet'
)

# 8. ТЕСТИРОВАНИЕ И ВЫПОЛНЕНИЕ
doc.add_heading('8. Результаты выполнения', level=1)

doc.add_paragraph(
    'Программа успешно:'
)

results = [
    'Компилируется без ошибок',
    'Выполняется без исключений (обрабатывает их правильно)',
    'Демонстрирует все три уровня иерархии Story→EventSequence→Event',
    'Выводит полный сценарий взаимодействия персонажей и предметов',
    'Использует все реализованные исключения в контролируемом виде',
    'Показывает корректность equals/hashCode через работу с объектами'
]

for result in results:
    doc.add_paragraph(result, style='List Bullet')

# 9. ВЫВОДЫ
doc.add_heading('9. Выводы', level=1)

doc.add_paragraph(
    'Реализованный проект демонстрирует:'
)

conclusions = [
    'Правильное применение наследования и полиморфизма',
    'Использование интерфейсов для описания поведения',
    'Корректную обработку исключений с механизмом throws',
    'Архитектурный подход к построению сложных сценариев',
    'Правильную реализацию equals/hashCode для всех объектов',
    'Использование современных языковых конструкций Java (records, generics)'
]

for conclusion in conclusions:
    doc.add_paragraph(conclusion, style='List Bullet')

# Сохранение
doc.save('d:\\itmo-lab\\Java\\lab3\\report.docx')
print("Отчет успешно создан: report.docx")
