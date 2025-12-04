#!/usr/bin/env python3
"""Create a docx report from markdown"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Отчёт: Java Lab 3 - OOP Model of Literary Text', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 1: Introduction
doc.add_heading('1. Введение', level=1)
p = doc.add_paragraph(
    'Проект Lab 3 реализует объектно-ориентированную модель русского литературного текста '
    '"История про Шутило" с использованием принципов проектирования ООП, паттернов проектирования '
    'и архитектурных подходов.'
)
p = doc.add_paragraph(
    'Основная цель: создать архитектуру, которая моделирует события (Events), их последовательности '
    '(EventSequences) и полную историю (Story), позволяя описать сложные сценарии взаимодействия '
    'персонажей с предметами.'
)

# Section 2: Architecture
doc.add_heading('2. Архитектура', level=1)
doc.add_heading('2.1 Иерархия Story → EventSequence → Event', level=2)

p = doc.add_paragraph(
    'Архитектура построена на трёхуровневой иерархии:'
)
doc.add_paragraph('Story', style='List Bullet')
doc.add_paragraph('EventSequence 1, EventSequence 2, ...', style='List Bullet 2')
doc.add_paragraph('Event 1, Event 2, ...', style='List Bullet 3')

doc.add_paragraph(
    'Story: Контейнер для всех последовательностей событий. Метод tell() выполняет все последовательности.'
)
doc.add_paragraph(
    'EventSequence: Логическая группа событий с условиями, временем и объектами. Метод simulate() исполняет все события.'
)
doc.add_paragraph(
    'Event: Отдельное событие с субъектами (люди), типом события, временем и объектом. Метод happen() выводит описание события.'
)

doc.add_heading('2.2 Ключевые компоненты', level=2)
doc.add_heading('Персонажи (Characters)', level=3)
doc.add_paragraph('Human (abstract): Базовый класс для всех персонажей', style='List Bullet')
doc.add_paragraph('Shutilo: Главный персонаж (спешит и путает вещи)', style='List Bullet')
doc.add_paragraph('Korzhik: Вспомогательный персонаж', style='List Bullet')
doc.add_paragraph('Swisstulkin: Вспомогательный персонаж', style='List Bullet')

doc.add_heading('Предметы (Items)', level=3)
doc.add_paragraph('Item (abstract): Базовый класс для всех предметов', style='List Bullet')
doc.add_paragraph('Cloth (abstract): Одежда', style='List Bullet')
doc.add_paragraph('Tie, Pants, Chulki, Jacket: Различные виды одежды', style='List Bullet 2')
doc.add_paragraph('Car: Автомобиль', style='List Bullet')
doc.add_paragraph('DriverLicense: Водительское удостоверение', style='List Bullet')
doc.add_paragraph('FlowerPot: Горшок с цветком (может разбиться)', style='List Bullet')

# Section 3: Exception Handling
doc.add_heading('3. Обработка Исключений', level=1)
doc.add_heading('3.1 Custom Exceptions', level=2)

doc.add_paragraph('ClothingMiswearException (checked)', style='List Bullet')
doc.add_paragraph('Выбрасывается при неправильном надевании одежды', style='List Bullet 2')

doc.add_paragraph('ItemLostException (checked)', style='List Bullet')
doc.add_paragraph('Выбрасывается при потере предмета', style='List Bullet 2')

doc.add_paragraph('InventoryFullException (unchecked)', style='List Bullet')
doc.add_paragraph('Выбрасывается при переполнении инвентаря', style='List Bullet 2')

doc.add_heading('3.2 Обработка в Main', level=2)
p = doc.add_paragraph(
    'В Main.java исключения обрабатываются через try-catch блоки, которые преобразуют исключения в fallback Events:'
)

# Section 4: OOP Principles
doc.add_heading('4. Реализация OOP принципов', level=1)

doc.add_heading('4.1 Инкапсуляция', level=2)
p = doc.add_paragraph('Все поля классов используют правильные модификаторы доступа:')
doc.add_paragraph('private: Приватные поля (не должны быть доступны извне)', style='List Bullet')
doc.add_paragraph('protected: Защищённые поля (доступны наследникам)', style='List Bullet')
doc.add_paragraph('public: Публичные поля (только Event.subject для удобства)', style='List Bullet')

doc.add_heading('4.2 Наследование', level=2)
p = doc.add_paragraph('Классы образуют правильную иерархию:')
doc.add_paragraph('Human ← Shutilo, Korzhik, Swisstulkin', style='List Bullet')
doc.add_paragraph('Item ← Cloth ← Tie, Pants, Chulki, Jacket', style='List Bullet')
doc.add_paragraph('Item ← Car, DriverLicense', style='List Bullet')

doc.add_heading('4.3 Полиморфизм', level=2)
p = doc.add_paragraph('Переопределены методы в наследниках:')
doc.add_paragraph('wear() в каждом классе одежды с специфичной логикой', style='List Bullet')
doc.add_paragraph('changeState() в каждом классе предмета', style='List Bullet')
doc.add_paragraph('showInventory() в каждом персонаже', style='List Bullet')

doc.add_heading('4.4 equals(), hashCode(), toString()', level=2)
p = doc.add_paragraph('Все 13 классов моделей реализуют:')
doc.add_paragraph('equals(): Использует hashCode() для сравнения', style='List Bullet')
doc.add_paragraph('hashCode(): Учитывает ВСЕ значимые поля (множитель 31)', style='List Bullet')
doc.add_paragraph('toString(): Выводит все свойства объекта', style='List Bullet')

# Section 5: Main Architecture
doc.add_heading('5. Архитектура Main.java', level=1)
doc.add_heading('5.1 Принцип Design', level=2)
p = doc.add_paragraph('Main.java содержит ТОЛЬКО конструирование сценария, без:')
doc.add_paragraph('System.out.println (выводу информации)', style='List Bullet')
doc.add_paragraph('Бизнес-логики (условных операторов, случайных выборов)', style='List Bullet')
doc.add_paragraph('Явной обработки ошибок (вывода ошибок в консоль)', style='List Bullet')

doc.add_heading('5.2 Структура Main.java', level=2)
p = doc.add_paragraph('Основная структура:')
doc.add_paragraph('1. Создание персонажей', style='List Number')
doc.add_paragraph('2. Создание предметов', style='List Number')
doc.add_paragraph('3. Построение истории (Story)', style='List Number')
doc.add_paragraph('4. Создание последовательностей с событиями', style='List Number')
doc.add_paragraph('5. Обработка исключений через fallback events', style='List Number')
doc.add_paragraph('6. Выполнение истории через tale.tell()', style='List Number')

# Section 6: Design Patterns
doc.add_heading('6. Паттерны проектирования', level=1)

doc.add_heading('6.1 Builder Pattern (метод chaining)', level=2)
p = doc.add_paragraph(
    'Event использует builder pattern для удобного создания с методами addSubject(), addTime(), addObject():'
)

doc.add_heading('6.2 Template Method Pattern', level=2)
p = doc.add_paragraph(
    'EventSequence.simulate() использует template method для выполнения цепи событий.'
)

doc.add_heading('6.3 Decorator Pattern', level=2)
p = doc.add_paragraph(
    'Inventory декорирует Item по расположению в InventoryCell.'
)

# Section 7: Results
doc.add_heading('7. Результаты', level=1)

doc.add_heading('7.1 Компиляция', level=2)
p = doc.add_paragraph('✅ Все файлы успешно компилируются без ошибок')

doc.add_heading('7.2 Архитектурные требования', level=2)
p = doc.add_paragraph('Все требования выполнены:')
doc.add_paragraph('Story содержит EventSequence', style='List Bullet')
doc.add_paragraph('EventSequence содержит Event', style='List Bullet')
doc.add_paragraph('Event.happen() выполняет вывод события', style='List Bullet')
doc.add_paragraph('Main не содержит println', style='List Bullet')
doc.add_paragraph('Main не содержит бизнес-логики', style='List Bullet')
doc.add_paragraph('Main только конструирует сценарий', style='List Bullet')

doc.add_heading('7.3 OOP требования', level=2)
p = doc.add_paragraph('Все OOP требования выполнены:')
doc.add_paragraph('Все 13 классов имеют правильные equals/hashCode/toString', style='List Bullet')
doc.add_paragraph('Все методы с исключениями содержат throws', style='List Bullet')
doc.add_paragraph('Правильное наследование и полиморфизм', style='List Bullet')
doc.add_paragraph('Инкапсуляция с private/protected/public', style='List Bullet')

# Section 8: Conclusion
doc.add_heading('8. Заключение', level=1)
p = doc.add_paragraph(
    'Проект успешно реализует объектно-ориентированную архитектуру для моделирования сложных сценариев с использованием:'
)
doc.add_paragraph('Трёхуровневой иерархии (Story → EventSequence → Event)', style='List Bullet')
doc.add_paragraph('Custom exceptions для различных ошибочных ситуаций', style='List Bullet')
doc.add_paragraph('Правильного применения ООП принципов', style='List Bullet')
doc.add_paragraph('Чистой архитектуры Main.java без бизнес-логики', style='List Bullet')
doc.add_paragraph('Полного использования функций Java (интерфейсы, наследование, полиморфизм)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Код готов к дальнейшему расширению и добавлению новых типов событий, персонажей и предметов.').bold = True

# Save document
doc.save(r'd:\itmo-lab\Java\lab3\report.docx')
print("Document saved successfully: d:\\itmo-lab\\Java\\lab3\\report.docx")
